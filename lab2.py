import docx
import mtk2
from docx.shared import Pt

def take_text(path):
    qwe = []
    for p in path.paragraphs:
        qwe.append(p.text)
        p.clear()
    print(qwe)

    return qwe

def split_text(bits, text):
    tre = []
    tline = ''
    for i in range(len(bits)):
        if i != len(bits) - 1:
            if bits[i] == '0' and bits[i+1] == '0':
                tline += text[i]
            elif bits[i] == '1' and bits[i+1] == '1':
                tline += text[i]
            elif bits[i] == '0' and bits[i + 1] == '1':
                tline += text[i]
                tre.append(tline)
                tline = ''
            elif bits[i] == '1' and bits[i + 1] == '0':
                tline += text[i]
                tre.append(tline)
                tline = ''
        else:
            if bits[i] == '0' and bits[i - 1] == '0':
                tline += text[i]
                tre.append(tline)
            elif bits[i] == '1' and bits[i - 1] == '1':
                tline += text[i]
                tre.append(tline)
            elif bits[i] == '0' and bits[i - 1] == '1':
                tline = text[i]
                tre.append(tline)
            elif bits[i] == '1' and bits[i - 1] == '0':
                tline = text[i]
                tre.append(tline)

    return tre

def run_get_scale(run, scale):
    rPr = run._r.get_or_add_rPr()
    for elem in rPr:
        if elem.tag.endswith('w'):
            rPr.remove(elem)

    w_elem = rPr.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}spacing')
    w_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', str(scale))
    rPr.append(w_elem)

my_text = 'Хороший товарищ, который правду в глаза говорит.'
my_encoding = 'Бодо'
bit_line = mtk2.MTK2_code(my_text)

doc = docx.Document('9.docx')

array_text = array_text_split_to_p = take_text(doc)


array_text_lengths = []
for i in array_text:
    rwqe = []
    for j in i.split('\n'):
        rwqe.append(len(j))
    array_text_lengths.append([len('|'.join(i.split('\n'))), len(i.split('\n')) - 1, rwqe])
print(array_text_lengths)

need_paragraphs = 0
for i in range(len(array_text_lengths)):
    need_paragraphs += array_text_lengths[i][0] - array_text_lengths[i][1]
    if need_paragraphs >= len(bit_line):
        need_paragraphs = i + 1
        break
print(need_paragraphs)

bit_line_split_to_p = []
q = 0
right = 0
#print(bit_line)

for i in range(need_paragraphs):
    if q == 0:
        left = 0
        right += array_text_lengths[i][0]
        bit_line_split_to_p.append(bit_line[left:right])
        q = right
    elif q != 0:
        left = q
        right += array_text_lengths[i][0]

        temp_l = bit_line[left:right]
        if len(temp_l) != 0:
            #print(temp_l)
            w = 0
            temp_a = []
            right_2 = 0
            for k in range(len(array_text_lengths[i][2])):
                left_2 = w
                right_2 += array_text_lengths[i][2][k]
                if len(temp_l[left_2:right_2]) == 0:
                    break
                temp_a.append(temp_l[left_2:right_2])
                w = right_2
            bit_line_split_to_p.append(temp_a[0])
        else:
            temp_l = bit_line[left:]

            qcount = 0
            for ee in range(len(array_text_lengths[i][2])):
                qcount += array_text_lengths[i][2][ee]
                if qcount >= len(temp_l):
                    qcount = ee + 1
                    break

            temp_a = []
            w = 0
            right_2 = 0
            for ee in range(qcount):
                left_2 = w
                right_2 += array_text_lengths[i][2][ee]
                if len(temp_l[left_2:right_2]) != 0:
                    temp_a.append(temp_l[left_2:right_2])
                else:
                    temp_a.append(temp_l[left_2:])
                w = right_2

            bit_line_split_to_p.append(temp_a[0])

        q = right
print(bit_line_split_to_p)
print(array_text_split_to_p)
print()
print('Вывод деления строки ')
counterqweqw = 0
parameter_num = 21
for i in range(len(array_text)):
    p = doc.paragraphs[i]
    if i == 0:
        temp_array = split_text(bit_line_split_to_p[i], array_text_split_to_p[i])
        print(temp_array)
        for j in range(len(temp_array)):
            p.add_run(temp_array[j])
            if counterqweqw % 2 == 1:
                run_get_scale(p.runs[j], parameter_num)
            counterqweqw += 1
    elif i < need_paragraphs and i > 0:
        # ['010110011111100011100', '10000011010100110100', '10110101011100010', '01101100100100000110']
        # ['Шла вчера я за водою,', 'А у нас ведро худое.', 'Из-за этого ведра', 'Я наплакалась вчера.']
        if len(bit_line_split_to_p[i]) == len(array_text_split_to_p[i]):
            giga_array = split_text(bit_line_split_to_p[i], array_text_split_to_p[i])
            print(giga_array)
            for k in range(len(giga_array)):
                p.add_run(giga_array[k])
                if counterqweqw % 2 == 1:
                    run_get_scale(p.runs[k], parameter_num)
                counterqweqw += 1


        # ['0011000000110000', '10110010011101010000', '101110']
        # ['Залепила я дыру.', 'Только воду наберу —', 'А она опять наружу', 'Так и льется по ведру.']
        elif len(bit_line_split_to_p[i]) != len(array_text_split_to_p[i]):
            count1 = 0
            count2 = len(array_text_split_to_p[i])
            if len(bit_line_split_to_p[i]) == len(array_text_split_to_p[i]):
                giga_array = split_text(bit_line_split_to_p[i], array_text_split_to_p[i])
                print(giga_array)
                for k in range(len(giga_array)):
                    p.add_run(giga_array[k])
                    if counterqweqw % 2 == 1:
                        run_get_scale(p.runs[k], parameter_num)
                    counterqweqw += 1

                count1 += 1
            elif len(bit_line_split_to_p[i]) != len(array_text_split_to_p[i]):
                bits_temp = bit_line_split_to_p[i]
                text_temp = array_text_split_to_p[i][:len(bits_temp)]
                giga_array = split_text(bits_temp, text_temp)
                ost_tt = array_text_split_to_p[i][len(bits_temp):]
                print(giga_array)
                for k in range(len(giga_array)):
                    p.add_run(giga_array[k])
                    if counterqweqw % 2 == 1:
                        run_get_scale(p.runs[k], parameter_num)
                    counterqweqw += 1

                p.add_run(ost_tt)
                #p.add_run('\n')
                count1 += 1
            revq = array_text_split_to_p[i][::-1]


    elif i >= need_paragraphs:
        p.add_run(array_text_split_to_p[i])

print('Добавление оставшегося текста...')
print('Завершено')

print('Изменение параметров...')

for i in range(len(doc.paragraphs)):
    p = doc.paragraphs[i]
    if i == 0:
        for run in p.runs:
            run.font.size = Pt(14)
            run.font.name = 'Bahnschrift Condensed'
    elif i != 0:
        for run in p.runs:
            if run.text == '\n':
                continue
            else:
                run.font.size = Pt(14)
                run.font.name = 'Bahnschrift Condensed'
print('Завершено изменение параметров')


doc.save('Итог.docx')
print('Текст сохранен в Итог.docx')

