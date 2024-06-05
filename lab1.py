import docx


doc = docx.Document('variant19.docx')


text = []
for paragraph in doc.paragraphs:
    text.append(paragraph.text)
#print('\n'.join(text))

code = ''
q = 0

e = False
if doc.paragraphs[0].runs[0].font.color.rgb != doc.paragraphs[0].runs[1].font.color.rgb:
    print('\nСпособ форматирования - изменение цвета символов.')
elif doc.paragraphs[0].runs[0].font.highlight_color != doc.paragraphs[0].runs[1].font.highlight_color:
    print('\nСпособ форматирования - изменение цвета фона.')
elif doc.paragraphs[0].runs[0].font.size.pt != doc.paragraphs[0].runs[1].font.size.pt:
    print('\nСпособ форматирования - изменение размера шрифта.')
else:
    print('\nСпособ форматирования - изменение масштаба шрифта или межсимвольного интервала.')
    e = True

for i in range(len(doc.paragraphs)):
    for j in range(len(doc.paragraphs[i].runs)):
        if doc.paragraphs[i].text != doc.paragraphs[i].runs[j].text:
            if not e and j != 0 and doc.paragraphs[i].runs[j].font.color.rgb == doc.paragraphs[i].runs[j - 1].font.color.rgb and doc.paragraphs[i].runs[j].font.highlight_color == doc.paragraphs[i].runs[j - 1].font.highlight_color and doc.paragraphs[i].runs[j].font.size.pt == doc.paragraphs[i].runs[j - 1].font.size.pt:
                if q == 0:
                    q = 1
                else:
                    q = 0
                code += str(q) * len(doc.paragraphs[i].runs[j].text)
                print(doc.paragraphs[i].runs[j].text, '-', str(q) * len(doc.paragraphs[i].runs[j].text),
                      doc.paragraphs[i].runs[j].font.color.rgb, doc.paragraphs[i].runs[j].font.highlight_color,
                      doc.paragraphs[i].runs[j].font.size.pt)

            else:
                code += str(q) * len(doc.paragraphs[i].runs[j].text)
                print(doc.paragraphs[i].runs[j].text, '-', str(q) * len(doc.paragraphs[i].runs[j].text),
                      doc.paragraphs[i].runs[j].font.color.rgb, doc.paragraphs[i].runs[j].font.highlight_color,
                      doc.paragraphs[i].runs[j].font.size.pt)

            if i != len(doc.paragraphs) - 1:
                if doc.paragraphs[i + 1].runs:
                    if not(j == len(doc.paragraphs[i].runs) - 1 and doc.paragraphs[i].runs[j].font.color.rgb == doc.paragraphs[i + 1].runs[0].font.color.rgb and doc.paragraphs[i].runs[j].font.highlight_color == doc.paragraphs[i + 1].runs[0].font.highlight_color and doc.paragraphs[i].runs[j].font.size.pt == doc.paragraphs[i + 1].runs[0].font.size.pt):
                        if q == 0:
                            q = 1
                        else:
                            q = 0
                else:
                    if not(j == len(doc.paragraphs[i].runs) - 1 and doc.paragraphs[i].runs[j].font.color.rgb == doc.paragraphs[i + 2].runs[0].font.color.rgb and doc.paragraphs[i].runs[j].font.highlight_color == doc.paragraphs[i + 2].runs[0].font.highlight_color and doc.paragraphs[i].runs[j].font.size.pt == doc.paragraphs[i + 2].runs[0].font.size.pt):
                        if q == 0:
                            q = 1
                        else:
                            q = 0
        #print(doc.paragraphs[i].runs[j].text, '-', str(q) * len(doc.paragraphs[i].runs[j].text), doc.paragraphs[i].runs[j].font.color.rgb, doc.paragraphs[i].runs[j].font.highlight_color, doc.paragraphs[i].runs[j].font.size.pt)


print('1 вариант сокрытого кода: ', code)
code2 = list(code)
for i in range(len(code2)):
    if code2[i] == '0':
        code2[i] = '1'
    else:
        code2[i] = '0'
code2 = ''.join(code2)
print('2 вариант сокрытого кода: ', code2)

ArrayBytesCod = ['00011', '11001', '01110', '01001', '00001', '01101', '11010', '10100', '00110', '01011',
                 '01111', '10010', '11100', '01100', '11000', '10110', '10111', '01010', '00101', '10000',
                 '00111', '11110', '10011', '11101', '10101', '10001']
ArrayLatinLetters = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M', 'N', 'O', 'P', 'Q',
                     'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z']
def decodingBodo(code):
    revealed_text = ''
    for i in range(0, len(code), 5):
        if code[i : i + 5] in ArrayBytesCod:
            revealed_text += ArrayLatinLetters[ArrayBytesCod.index(code[i : i + 5])]
    return revealed_text

print('\nКод Бодо (МТК-2)')
revealed_text = decodingBodo(code)
print(revealed_text)
revealed_text = decodingBodo(code2)
print(revealed_text)

ArrayBytesCodKoi8 = ['11000000','11000001','11000010', '11000011','11000100', '11000101','11000110', '11000111',
                     '11001000','11001001','11001010', '11001011','11001100', '11001101','11001110', '11001111',
                     '11010000','11010001','11010010', '11010011','11010100', '11010101','11010110', '11010111',
                     '11011000','11011001','11011010', '11011011','11011100', '11011101','11011110', '11011111',
                     '11100000', '11100001', '11100010', '11100011', '11100100', '11100101', '11100110', '11100111',
                     '11101000', '11101001', '11101010', '11101011', '11101100', '11101101', '11101110', '11101111',
                     '11110000', '11110001', '11110010', '11110011', '11110100', '11110101', '11110110', '11110111',
                     '11111000', '11111001', '11111010', '11111011', '11111100', '11111101', '11111110', '11111111', '00100000', '00101001']
ArraySymbolKoi8 = ['ю', 'а', 'б', 'ц', 'д', 'е', 'ф', 'г', 'х', 'и', 'й', 'к', 'л', 'м', 'н', 'о',
                   'п', 'я', 'р', 'с', 'т', 'у', 'ж', 'в', 'ь', 'ы', 'з', 'ш', 'э', 'щ', 'ч', 'ъ',
                   'Ю', 'А', 'Б', 'Ц', 'Д', 'Е', 'Ф', 'Г', 'Х', 'И', 'Й', 'К', 'Л', 'М', 'Н', 'О',
                   'П', 'Я', 'Р', 'С', 'Т', 'У', 'Ж', 'В', 'Ь', 'Ы', 'З', 'Ш', 'Э', 'Щ', 'Ч', 'Ъ', ' ', '-']


def decodingKoi(code):
    revealed_text = ''
    for i in range(0, len(code), 8):
        if code[i : i + 8] in ArrayBytesCodKoi8:
            revealed_text += ArraySymbolKoi8[ArrayBytesCodKoi8.index(code[i : i + 8])]
    return revealed_text

print('\nКОИ-8R')
revealed_text = decodingKoi(code)
print(revealed_text)
revealed_text = decodingKoi(code2)
print(revealed_text)

ArrayBytesCodCP866 = ['10000000','10000001','10000010', '10000011','10000100', '10000101','10000110', '10000111',
                     '10001000','10001001','10001010', '10001011','10001100', '10001101','10001110', '10001111',
                     '10010000','10010001','10010010', '10010011','10010100', '10010101','10010110', '10010111',
                     '10011000','10011001','10011010', '10011011','10011100', '10011101','10011110', '10011111',
                      '10100000', '10100001', '10100010', '10100011', '10100100', '10100101', '10100110', '10100111',
                      '10101000', '10101001', '10101010', '10101011', '10101100', '10101101', '10101110', '10101111',
                      '11100000', '11100001', '11100010', '11100011', '11100100', '11100101', '11100110', '11100111',
                      '11101000', '11101001', '11101010', '11101011', '11101100', '11101101', '11101110', '11101111', '00100000']
ArraySymbolCP866 = ['А', 'Б', 'В', 'Г', 'Д', 'Е', 'Ж', 'З', 'И', 'Й', 'К', 'Л', 'М', 'Н', 'О', 'П',
                     'Р', 'С', 'Т', 'У', 'Ф', 'Х', 'Ц', 'Ч', 'Ш', 'Щ', 'Ъ', 'Ы', 'Ь', 'Э', 'Ю', 'Я',
                     'а', 'б', 'в', 'г', 'д', 'е', 'ж', 'з', 'и', 'й', 'к', 'л', 'м', 'н', 'о', 'п',
                     'р', 'с', 'т', 'у', 'ф', 'х', 'ц', 'ч', 'ш', 'щ', 'ъ', 'ы', 'ь', 'э', 'ю', 'я', ' ']

def decodingCP866(code):
    revealed_text = ''
    for i in range(0, len(code), 8):
        #print(code[i: i + 8])
        if code[i : i + 8] in ArrayBytesCodCP866:
            revealed_text += ArraySymbolCP866[ArrayBytesCodCP866.index(code[i: i + 8])]
            #print(ArraySymbolCP866[ArrayBytesCodCP866.index(code[i: i + 8])], code[i: i + 8])
    return revealed_text

print('\nCP866')
revealed_text = decodingCP866(code)
print(revealed_text)
revealed_text = decodingCP866(code2)
print(revealed_text)

ArrayBytesCodWin = ['11000000','11000001','11000010', '11000011','11000100', '11000101','11000110', '11000111',
                     '11001000','11001001','11001010', '11001011','11001100', '11001101','11001110', '11001111',
                     '11010000','11010001','11010010', '11010011','11010100', '11010101','11010110', '11010111',
                     '11011000','11011001','11011010', '11011011','11011100', '11011101','11011110', '11011111',
                    '11100000', '11100001', '11100010', '11100011', '11100100', '11100101', '11100110', '11100111',
                    '11101000', '11101001', '11101010', '11101011', '11101100', '11101101', '11101110', '11101111',
                    '11110000', '11110001', '11110010', '11110011', '11110100', '11110101', '11110110', '11110111',
                    '11111000', '11111001', '11111010', '11111011', '11111100', '11111101', '11111110', '11111111', '00100000']
ArraySymbolWin = ['А', 'Б', 'В', 'Г', 'Д', 'Е', 'Ж', 'З', 'И', 'Й', 'К', 'Л', 'М', 'Н', 'О', 'П',
                  'Р', 'С', 'Т', 'У', 'Ф', 'Х', 'Ц', 'Ч', 'Ш', 'Щ', 'Ъ', 'Ы', 'Ь', 'Э', 'Ю', 'Я',
                  'а', 'б', 'в', 'г', 'д', 'е', 'ж', 'з', 'и', 'й', 'к', 'л', 'м', 'н', 'о', 'п',
                  'р', 'с', 'т', 'у', 'ф', 'х', 'ц', 'ч', 'ш', 'щ', 'ъ', 'ы', 'ь', 'э', 'ю', 'я', ' ']

def decodingWin(code):
    revealed_text = ''
    for i in range(0, len(code), 8):
        if code[i : i + 8] in ArrayBytesCodWin:
            revealed_text += ArraySymbolWin[ArrayBytesCodWin.index(code[i: i + 8])]
    return revealed_text

print('\nWindows-1251')
revealed_text = decodingWin(code)
print(revealed_text)
c1 = ''
for i in range(0, len(code), 8):
    c1 += code[i : i + 8] + ' '
print(c1)
revealed_text = decodingWin(code2)
print(revealed_text)
c1 = ''
for i in range(0, len(code2), 8):
    c1 += code2[i : i + 8] + ' '
print(c1)
